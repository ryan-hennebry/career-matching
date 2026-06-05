from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "output/cv/ryan-hennebry-cv-paraglide-highlighted.docx"

BLACK = RGBColor(0x11, 0x11, 0x11)
BLUE = RGBColor(0x1F, 0x5A, 0xA6)
GREY = RGBColor(0x4B, 0x4B, 0x4B)


def set_font(run, name="Barlow", size=10.8, bold=False, italic=False, color=BLACK, underline=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.underline = underline


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_border_bottom(paragraph, color="333333", size="8", space="1"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_bullet_paragraph(doc, label, text, color=BLACK):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.02)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.05)
    fmt.first_line_indent = Inches(0)
    bullet = p.add_run("● ")
    set_font(bullet, size=10.1, color=color)
    lead = p.add_run(f"{label}: ")
    set_font(lead, size=10.1, bold=True, color=color)
    body = p.add_run(text)
    set_font(body, size=10.1, color=color)
    return p


def add_role_line(doc, title, dates, location=None, color=BLACK):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=0, line=1.0)
    title_run = p.add_run(title)
    set_font(title_run, size=11.1, bold=True, color=color)
    sep = p.add_run(" · ")
    set_font(sep, size=11.1, color=color)
    date_run = p.add_run(dates)
    set_font(date_run, size=11.1, italic=True, color=color)
    if location:
        loc_sep = p.add_run(" ")
        set_font(loc_sep, size=11.1, color=color)
        loc_run = p.add_run(location)
        set_font(loc_run, size=11.1, italic=True, color=color)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=9, after=3, line=1.0)
    run = p.add_run(text)
    set_font(run, size=10.3, bold=True)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    normal = doc.styles["Normal"]
    normal.font.name = "Barlow"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Barlow")
    normal.font.size = Pt(10.3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    r = p.add_run("Ryan Hennebry")
    set_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    r = p.add_run("Cambridge, UK · Open to London (hybrid)")
    set_font(r, size=10, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=5)
    for i, label in enumerate(["Email", "LinkedIn", "GitHub"]):
        run = p.add_run(label)
        set_font(run, size=10, color=GREY, underline=True)
        if i < 2:
            sep = p.add_run(" · ")
            set_font(sep, size=10, color=GREY)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=3, line=1.08)
    r1 = p.add_run(
        "Founder-side startup operator and first employee at Minima, trusted to own high-priority work as the company scaled from seed to Series A across fundraising, community, GTM, and marketing. "
    )
    set_font(r1, size=11, color=BLUE)
    r2 = p.add_run(
        "Recently built AI agents that automate competitor intelligence, growth experiments, and career matching."
    )
    set_font(r2, size=11, color=BLUE)

    add_section_heading(doc, "EXPERIENCE")

    add_role_line(doc, "AI Workflow Automation Projects", "Feb 2025 – Present", "(Cambridge, UK/Remote)", color=BLUE)
    add_bullet_paragraph(
        doc,
        "Built",
        "a competitor-intelligence agent that researches markets and generates concise briefing PDFs with recommended actions.",
        color=BLUE,
    )
    add_bullet_paragraph(
        doc,
        "Built",
        "a growth-experiments agent that turns campaign performance into a prioritized next test with draft-ready copy.",
        color=BLUE,
    )
    add_bullet_paragraph(
        doc,
        "Built",
        "a career-matching agent that parses CVs, ranks live roles by fit, and generates tailored application briefs.",
        color=BLUE,
    )

    add_role_line(doc, "Minima", "Mar 2019 – Feb 2025", "(London, UK/Remote)")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=1, line=1.0)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.05)
    run = p.add_run(
        "First employee; progressed from Founder’s Associate to Community Manager to Marketing Manager as company scaled from seed to Series A."
    )
    set_font(run, size=10, color=BLUE)

    add_role_line(doc, "Marketing Manager", "Mar 2022 – Feb 2025")
    add_bullet_paragraph(
        doc,
        "Product Marketing",
        "Led the go-to-market for a community-led presale that delivered $4.5m on a $300k budget (15x ROI).",
        color=BLUE,
    )
    add_bullet_paragraph(
        doc,
        "Event Marketing",
        "Secured event sponsorships and keynote slots that connected founders with investors, contributing to Minima’s $6.5m Series A raise.",
    )
    add_bullet_paragraph(
        doc,
        "Email Marketing",
        "Generated 5k presale conversions from a 50k community email list through segmentation, A/B testing, and targeted messaging.",
        color=BLUE,
    )

    add_role_line(doc, "Community Manager", "Apr 2020 – Mar 2022")
    add_bullet_paragraph(
        doc,
        "Community Growth",
        "Delivered Minima’s Incentive Program, onboarding 50k node-runners and scaling to 75k nodes on the network.",
    )
    add_bullet_paragraph(
        doc,
        "Community Management",
        "Grew Discord and Telegram to 100k members through an ambassador programme and hackathon series.",
    )

    add_role_line(doc, "Founder’s Associate", "Mar 2019 – Apr 2020")
    add_bullet_paragraph(
        doc,
        "Fundraising",
        "Supported the $2.5m seed round by handling investor enquiries, pitch materials, and updates.",
        color=BLUE,
    )
    add_bullet_paragraph(
        doc,
        "Execution",
        "Worked directly with founders on high-priority work across community, operations, and early marketing.",
    )
    add_bullet_paragraph(
        doc,
        "Ownership",
        "Managed multiple workstreams and brought structure to early-stage chaos without slowing execution.",
    )

    add_role_line(doc, "Content Marketer, Blockchain Board of Derivatives", "Oct 2018 – Mar 2019", "(Cambridge, UK)")
    add_bullet_paragraph(
        doc,
        "User Growth",
        "Acquired 20k exchange users through technical content creation and distribution.",
    )

    add_role_line(doc, "Data Analyst, Amazon Alexa", "Oct 2017 – Oct 2018", "(Cambridge, UK)")
    add_bullet_paragraph(
        doc,
        "Quality Improvement",
        "Improved Alexa’s success rate to 98% on key queries via failure analysis.",
    )

    add_section_heading(doc, "EDUCATION")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.0)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.05)
    bullet = p.add_run("● ")
    set_font(bullet, size=10.1)
    uni = p.add_run("University of Leeds")
    set_font(uni, size=10.1, bold=True)
    body = p.add_run(" · 2017 · Leeds, UK · BA Management, Upper Second Class Honours (2:1)")
    set_font(body, size=10.1)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
